import networkx as nx
import requests
from collections import OrderedDict
import json
import os
import cairosvg
from collections import defaultdict
from docx import Document

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

method_id_map = {}
edges_dict = {}
call_dict = OrderedDict()
branches_num_dict = {}


document = Document()

def load_branches_num_data():
    file_path = "AMUST覆盖数据20230613.json"
    with open(file_path) as file:
        data = json.load(file)

    all_data = data['AllData'][1]
    method_cov = all_data['MethodCov']

    branches_num_dict = {}

    for method in method_cov:
        class_name = method['ClassName']
        method_name = method['MethodName']
        coverage = method['Branch']
        branches_num = method['cdcSumblock']
        # Create the key as class_name + ":" + method_name + ":" + coverage
        key = method_name + ":" + coverage
        # Assign branches_num to the corresponding key in the branches_num_dict
        branches_num_dict[key] = branches_num

    for key, value in branches_num_dict.items():
        print("num map ", key, value)

    key_to_check = "MethodName:Branch"
    if key_to_check in branches_num_dict:
        print("Key exists in branches_num_dict")
    else:
        print("Key does not exist in branches_num_dict")

def preorder_traversal(tree, node):
    print("# node:", node)
    print("##      :", edges_dict.get(node))
    index = method_id_map[node].find('-') - 1
    print("##      :", method_id_map[node][0:index])
    method_name = method_id_map[node][0:index]
    coverage = method_id_map[node][index+1:]
    key = method_name + ":" + coverage
    print("##key      :", key)
    branches_num = branches_num_dict.get(key)
    if branches_num:
        branch_coverage = method_id_map[node]
        headling_info = method_id_map[node][0:_] + " 分支数: " + branches_num + " 分支覆盖率: " + branch_coverage
        document.add_heading(headling_info, 0)
        pngfname = get_picture_by_id(node)
        document.add_picture(pngfname)
    else:
        print("Branches number not found for key:", key)

    for child in tree.successors(node):
        preorder_traversal(tree, child)

# The rest of your code...

def main():
    method_map()
    load_branches_num_data()

    # The rest of your code...

if __name__ == "__main__":
    main()

