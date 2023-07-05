import docx
from docx import Document
doc = Document('test.docx')
split_index = 1

document1 = Document()
document2 = Document()

for i, paragraph in enumerate(doc.paragraphs):
    if i < split_index:
        document1.add_paragraph(paragraph.text)
    else:
        document2.add_paragraph(paragraph.text)

document1.save('1.docx')
document2.save('2.docx')

